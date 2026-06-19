import json
from io import BytesIO
from uuid import uuid4

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from fastapi import APIRouter, HTTPException, Response, status
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.storage import VOCABULARY_FILE, ensure_storage_directories

router = APIRouter(prefix="/api/vocabulary", tags=["vocabulary"])


class VocabularyItem(BaseModel):
    id: str
    word: str
    lemma: str
    part_of_speech: str
    plural: str
    translation: str
    example_sentence: str
    source_page: int | None = None


class VocabularyCreateRequest(BaseModel):
    word: str = ""
    lemma: str
    part_of_speech: str
    plural: str = ""
    translation: str
    example_sentence: str
    source_page: int | None = None


class VocabularyAddResponse(BaseModel):
    item: VocabularyItem
    created: bool
    message: str


class VocabularyDeleteResponse(BaseModel):
    deleted: bool
    id: str


def _load_vocabulary() -> list[dict[str, object]]:
    ensure_storage_directories()

    try:
        data = json.loads(VOCABULARY_FILE.read_text(encoding="utf-8"))
    except json.JSONDecodeError as exc:
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Vocabulary storage is corrupted.",
        ) from exc

    if not isinstance(data, list):
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail="Vocabulary storage must contain a list.",
        )

    return data


def _save_vocabulary(items: list[dict[str, object]]) -> None:
    ensure_storage_directories()
    VOCABULARY_FILE.write_text(
        json.dumps(items, ensure_ascii=False, indent=2) + "\n",
        encoding="utf-8",
    )


def _required_text(value: str, field_name: str) -> str:
    cleaned = value.strip()
    if not cleaned:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail=f"{field_name} cannot be empty.",
        )
    return cleaned


def _normalize_duplicate_key(item: dict[str, object]) -> tuple[str, str]:
    lemma = str(item.get("lemma", "")).strip().casefold()
    part_of_speech = str(item.get("part_of_speech", "")).strip().casefold()
    return lemma, part_of_speech


def _cell_text(item: dict[str, object], key: str) -> str:
    value = str(item.get(key, "")).strip()
    return value or "-"


def _set_cell_shading(cell, fill: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    cell._tc.get_or_add_tcPr().append(shading)


def _set_cell_margins(cell, top: int = 80, bottom: int = 80, start: int = 120, end: int = 120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)

    for margin_name, value in {
        "top": top,
        "bottom": bottom,
        "start": start,
        "end": end,
    }.items():
        node = tc_mar.find(qn(f"w:{margin_name}"))
        if node is None:
            node = OxmlElement(f"w:{margin_name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def _set_table_geometry(table, column_widths_dxa: list[int]) -> None:
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(column_widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_indent = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_indent is None:
        tbl_indent = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_indent)
    tbl_indent.set(qn("w:w"), "120")
    tbl_indent.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    for grid_col in list(grid):
        grid.remove(grid_col)
    for width in column_widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)

    for row in table.rows:
        for cell, width in zip(row.cells, column_widths_dxa):
            _set_cell_width(cell, width)
            _set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def _format_paragraph(paragraph, font_size: int = 10, bold: bool = False) -> None:
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.15
    for run in paragraph.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(font_size)
        run.font.bold = bold


def _build_vocabulary_docx(items: list[dict[str, object]]) -> BytesIO:
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = document.add_paragraph()
    title.paragraph_format.space_after = Pt(10)
    title_run = title.add_run("Vocabulary Notebook")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(16)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(46, 116, 181)

    meta = document.add_paragraph()
    meta_run = meta.add_run(f"{len(items)} saved word{'s' if len(items) != 1 else ''}")
    meta_run.font.name = "Calibri"
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(90, 90, 90)

    headers = ["Lemma", "POS", "Plural", "Chinese Translation", "Example Sentence"]
    rows = max(len(items), 1) + 1
    table = document.add_table(rows=rows, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    table.autofit = False
    widths = [1550, 1050, 1450, 2200, 3110]

    for index, header in enumerate(headers):
        cell = table.cell(0, index)
        cell.text = header
        _set_cell_shading(cell, "E8EEF5")
        for paragraph in cell.paragraphs:
            _format_paragraph(paragraph, font_size=10, bold=True)

    if items:
        for row_index, item in enumerate(items, start=1):
            values = [
                _cell_text(item, "lemma"),
                _cell_text(item, "part_of_speech"),
                _cell_text(item, "plural"),
                _cell_text(item, "translation"),
                _cell_text(item, "example_sentence"),
            ]
            for column_index, value in enumerate(values):
                cell = table.cell(row_index, column_index)
                cell.text = value
                for paragraph in cell.paragraphs:
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                        if column_index in {1, 2}
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    _format_paragraph(paragraph, font_size=10)
    else:
        cell = table.cell(1, 0)
        cell.text = "No saved vocabulary."
        for paragraph in cell.paragraphs:
            _format_paragraph(paragraph, font_size=10)
        for column_index in range(1, len(headers)):
            table.cell(1, column_index).text = "-"

    _set_table_geometry(table, widths)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer


@router.get("", response_model=list[VocabularyItem])
def list_vocabulary() -> list[dict[str, object]]:
    return _load_vocabulary()


@router.get("/export")
def export_vocabulary() -> StreamingResponse:
    items = _load_vocabulary()
    docx_buffer = _build_vocabulary_docx(items)

    return StreamingResponse(
        docx_buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={
            "Content-Disposition": 'attachment; filename="vocabulary-notebook.docx"',
        },
    )


@router.post("", response_model=VocabularyAddResponse)
def add_vocabulary_item(
    request: VocabularyCreateRequest,
    response: Response,
) -> VocabularyAddResponse:
    lemma = _required_text(request.lemma, "Lemma")
    part_of_speech = _required_text(request.part_of_speech, "Part of speech")
    translation = _required_text(request.translation, "Translation")
    example_sentence = _required_text(request.example_sentence, "Example sentence")
    word = request.word.strip() or lemma

    items = _load_vocabulary()
    incoming_key = (lemma.casefold(), part_of_speech.casefold())

    for item in items:
        if _normalize_duplicate_key(item) == incoming_key:
            return VocabularyAddResponse(
                item=VocabularyItem.model_validate(item),
                created=False,
                message="Vocabulary item already exists.",
            )

    item = VocabularyItem(
        id=uuid4().hex,
        word=word,
        lemma=lemma,
        part_of_speech=part_of_speech,
        plural=request.plural.strip(),
        translation=translation,
        example_sentence=example_sentence,
        source_page=request.source_page,
    )
    items.append(item.model_dump())
    _save_vocabulary(items)
    response.status_code = status.HTTP_201_CREATED

    return VocabularyAddResponse(
        item=item,
        created=True,
        message="Vocabulary item saved.",
    )


@router.delete("/{item_id}", response_model=VocabularyDeleteResponse)
def delete_vocabulary_item(
    item_id: str,
) -> VocabularyDeleteResponse:
    items = _load_vocabulary()
    remaining = [item for item in items if str(item.get("id")) != item_id]

    if len(remaining) == len(items):
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="Vocabulary item was not found.",
        )
    _save_vocabulary(remaining)
    return VocabularyDeleteResponse(deleted=True, id=item_id)
